# -*- coding: utf-8 -*-
"""Genera la documentación TÉCNICA (Word) de la aplicación."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WHALE = RGBColor(0xF6, 0x18, 0x6A)
GRIS = RGBColor(0x55, 0x55, 0x55)

doc = Document()
base = doc.styles['Normal']
base.font.name = 'Calibri'
base.font.size = Pt(11)


def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = WHALE
    return p


def h2(t):
    return doc.add_heading(t, level=2)


def p(texto, negrita=False):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.bold = negrita
    return par


def code(texto):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return par


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def nums(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def tabla(encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = 'Light Grid Accent 1'
    for i, e in enumerate(encabezados):
        c = t.rows[0].cells[i]
        c.text = e
        for par in c.paragraphs:
            for run in par.runs:
                run.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = str(val)
    doc.add_paragraph()
    return t


# ----------------------------------------------------------------- PORTADA
tit = doc.add_paragraph()
tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tit.add_run("Locking System")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = WHALE
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Documentación técnica")
rs.font.size = Pt(16)
rs.font.color.rgb = GRIS
d2 = doc.add_paragraph()
d2.alignment = WD_ALIGN_PARAGRAPH.CENTER
d2.add_run("Aplicación Django desplegada en Render · https://whaletv.onrender.com")
doc.add_paragraph()

# ----------------------------------------------------------------- 1
h1("1. Visión general de la arquitectura")
p(
    "Locking System es una aplicación web monolítica construida con Django (patrón "
    "MVT: Modelo–Vista–Template). Su función es administrar el estado de bloqueo de "
    "televisores vendidos a cuotas y reflejar ese estado en un portal externo de "
    "bloqueo (WhaleTV Lock Management, lockservice.whaletv.com), que no expone una API "
    "utilizable para esto, por lo que la integración se hace mediante automatización de "
    "navegador (Selenium)."
)
p("Componentes principales:", negrita=True)
bullets([
    "Servidor web Django (gunicorn) que sirve las páginas y la lógica de negocio.",
    "Base de datos PostgreSQL como almacén persistente.",
    "Módulo de automatización (Selenium + Chromium headless) que opera el portal externo como lo haría un humano.",
    "Procesos en segundo plano (hilos de Python) para las operaciones masivas (sincronización y validación), con una página de progreso que consulta el estado por polling.",
    "Frontend con plantillas Django + Tailwind CSS (CDN) y JavaScript plano; estáticos servidos por WhiteNoise.",
    "Despliegue en Render usando un contenedor Docker (necesario para incluir Chromium).",
])
p("Flujo de datos resumido:", negrita=True)
code(
    "Usuario → Django (vistas) → PostgreSQL (estado local)\n"
    "Operación masiva → hilo en segundo plano → Selenium/Chromium → Portal externo\n"
    "Portal externo → (lectura) → Django actualiza SyncJobItem → Página de progreso (polling)"
)

# ----------------------------------------------------------------- 2
h1("2. Stack tecnológico y dependencias")
tabla(
    ["Paquete", "Versión", "Para qué se usa"],
    [
        ["Python", "3.13", "Lenguaje base (imagen python:3.13-slim en Docker)."],
        ["Django", "6.0.5", "Framework web (ORM, vistas, plantillas, admin, auth)."],
        ["psycopg[binary]", "3.3.4", "Driver de PostgreSQL."],
        ["dj-database-url", "3.1.2", "Parseo de DATABASE_URL a la config de Django."],
        ["gunicorn", "23.0.0", "Servidor WSGI de producción."],
        ["whitenoise", "6.12.0", "Servir archivos estáticos sin servidor aparte."],
        ["selenium", "4.40.0", "Automatización del navegador para operar el portal externo."],
        ["pandas", "2.2.3", "Lectura de archivos Excel/CSV en las importaciones."],
        ["openpyxl", "3.1.5", "Generación de archivos Excel en las exportaciones."],
        ["tzdata", "2025.2", "Zonas horarias (TIME_ZONE = America/Bogota)."],
        ["Tailwind CSS", "CDN", "Estilos del frontend (cargado desde CDN, sin build)."],
    ],
)
p(
    "No se usan dependencias externas para leer el archivo .env: hay un cargador propio "
    "(cargar_env en core/settings.py) que vuelca las variables a os.environ sin "
    "sobreescribir las ya definidas en el entorno real."
)

# ----------------------------------------------------------------- 3
h1("3. Estructura del proyecto")
code(
    "core/         Configuración del proyecto Django (settings, urls, wsgi).\n"
    "users/        App de usuarios: modelo de usuario personalizado (login por correo).\n"
    "whaletv/      App principal: modelos, vistas, urls, plantillas, portal_sync (Selenium).\n"
    "  models.py       Televisor, Bloqueo, SyncJob, SyncJobItem, RegistroSync, PinCodeGenerado.\n"
    "  views.py        Vistas (CRUD, importar/exportar, sincronizar, validar, progreso).\n"
    "  urls.py         Rutas de la app.\n"
    "  forms.py        TelevisorForm, BloqueoForm.\n"
    "  admin.py        Registro en el admin de Django.\n"
    "  portal_sync.py  Automatización Selenium del portal externo.\n"
    "  templates/whaletv/  Plantillas HTML (base.html y cada pantalla).\n"
    "  migrations/     Migraciones de base de datos.\n"
    "Dockerfile, render.yaml, Procfile, build.sh   Despliegue.\n"
    "requirements.txt  Dependencias Python.\n"
    "manage.py         Utilidad de línea de comandos de Django."
)
p(
    "Rutas registradas en core/urls.py: el panel de administración de Django en "
    "/admin/ y todas las rutas de la aplicación incluidas desde whaletv.urls en la raíz."
)

# ----------------------------------------------------------------- 4
h1("4. Modelo de datos")
p(
    "Todos los modelos de negocio viven en whaletv/models.py. El estado de bloqueo de un "
    "televisor NO se guarda como un dato suelto: se deriva de su historial de bloqueos."
)

h2("4.1 Televisor")
bullets([
    "mac_address (CharField 50, etiqueta 'Dirección MAC'): identificador del equipo.",
    "serial_number (CharField 50, 'Número de serie').",
    "numero_credito (CharField 60, 'Número de crédito'): solo dígitos, validado con RegexValidator ^\\d{1,60}$. Se guarda como texto porque 60 dígitos no caben en un entero.",
    "lock_status (BooleanField): estado actual (True = Bloqueado). Es un valor DERIVADO/persistido.",
    "created_at (DateTimeField, auto_now_add).",
])
p("Métodos y propiedades relevantes:", negrita=True)
bullets([
    "calcular_estado(): lee el último Bloqueo (ordenado por -created_at) y fija lock_status con su estado. Devuelve si cambió.",
    "actualizar_lock(): recalcula y persiste lock_status con un UPDATE directo (sin reentrar a save()).",
    "save(): si el objeto ya existe, recalcula el estado antes de guardar.",
    "refrescar_todos() (classmethod): recalcula el estado de todos los televisores.",
    "fecha_sincronizar (property): fecha que se empuja al portal. Bloqueado → hoy − 30 días; Desbloqueado → hoy + 30 días (constante DIAS_DESFASE = 30).",
    "ultimo_bloqueo (property): el Bloqueo más reciente.",
])

h2("4.2 Bloqueo")
p(
    "Historial de estados de un televisor. Cada importación de estado o cambio manual "
    "agrega un registro; el más reciente define el lock_status del Televisor."
)
bullets([
    "televisor (ForeignKey → Televisor, related_name='bloqueos', on_delete=CASCADE).",
    "serial_number, mac_address (copias para la bitácora).",
    "estado (BooleanField): True = Bloqueado.",
    "created_at (auto_now_add). Meta.ordering = ['-created_at'].",
    "save() y delete() llaman a televisor.actualizar_lock() para mantener el estado coherente.",
])

h2("4.3 SyncJob y SyncJobItem")
p(
    "Representan un trabajo masivo (en segundo plano) y el resultado por televisor. Se "
    "usan tanto para Sincronización como para Validación."
)
p("SyncJob:", negrita=True)
bullets([
    "tipo: 'sincronizacion' o 'validacion'.",
    "estado: pendiente / corriendo / terminado / error / cancelado (ACTIVOS = pendiente, corriendo).",
    "usuario (FK SET_NULL), usuario_email, total, workers, error, creado, actualizado, terminado.",
])
p("SyncJobItem:", negrita=True)
bullets([
    "job (FK → SyncJob, related_name='items'), televisor (FK SET_NULL), mac.",
    "estado: pendiente / ok / error.",
    "aplicado (BooleanField): si la sincronización cambió el portal.",
    "remoto_bloqueado, local_bloqueado, coincide: usados en validación (dry-run) para guardar el estado del portal, el de la app y si coinciden.",
    "mensaje (TextField): detalle o error.",
])

h2("4.4 RegistroSync (bitácora)")
bullets([
    "Quién sincronizó (usuario FK SET_NULL + usuario_email), televisor (FK SET_NULL) y copias planas (nombre_persona, mac_address) por si se borra el TV.",
    "lock_status (estado en el momento), aplicado, tipo ('individual' / 'masivo'), creado.",
    "Método de clase registrar(usuario, televisor, aplicado, tipo) que toma una 'foto' del estado.",
])

h2("4.5 PinCodeGenerado (bitácora de desbloqueo manual)")
bullets([
    "usuario (FK SET_NULL) + usuario_email, mac_address, passcode (Código de Acceso), pin_code (Código Pin), creado.",
])

# ----------------------------------------------------------------- 5
h1("5. Autenticación y usuarios")
p(
    "El proyecto usa un modelo de usuario personalizado (AUTH_USER_MODEL = 'users.User') "
    "definido en users/models.py sobre AbstractBaseUser + PermissionsMixin."
)
bullets([
    "El identificador de inicio de sesión es el correo (USERNAME_FIELD = 'email', único). REQUIRED_FIELDS = ['nombre'].",
    "Campos: email, nombre, is_active, is_staff, date_joined. Sin username.",
    "UserManager con create_user / create_superuser.",
    "El login se hace con authenticate(username=email, password=...) en la vista login_view; las vistas internas usan el decorador @login_required.",
    "Redirecciones: LOGIN_URL='login', LOGIN_REDIRECT_URL='home', LOGOUT_REDIRECT_URL='login'.",
])

# ----------------------------------------------------------------- 6
h1("6. Capa web: vistas, plantillas y frontend")
p(
    "Las vistas son funciones (function-based views) en whaletv/views.py. Cubren el CRUD "
    "de televisores, el cambio de estado (Bloqueo), las importaciones y exportaciones, "
    "la sincronización/validación (individual y masiva) y las bitácoras."
)
bullets([
    "Renderizado del lado del servidor con plantillas Django; base.html define el layout (sidebar, breadcrumb) y cada pantalla lo extiende.",
    "Estilos con Tailwind CSS cargado por CDN (sin paso de build); clases utilitarias y unos pocos componentes (.btn, .card, .pill, .th, .td) definidos en base.html.",
    "JavaScript plano (sin framework) para: la página de progreso (polling con fetch), el modal de generación de Código Pin y los toggles de previsualización.",
    "Mensajes al usuario con el framework de messages de Django.",
    "Archivos estáticos servidos por WhiteNoise con almacenamiento comprimido y con manifiesto (CompressedManifestStaticFilesStorage); se generan con collectstatic.",
])

# ----------------------------------------------------------------- 7
h1("7. Lógica de negocio clave")
h2("7.1 Estado derivado y fecha del portal")
p(
    "El lock_status de un Televisor es el estado de su último Bloqueo. La fecha que se "
    "envía al portal (fecha_sincronizar) se calcula sola: 30 días hacia atrás si está "
    "bloqueado (fecha vencida → el portal lo mantiene bloqueado) y 30 días hacia "
    "adelante si está desbloqueado (fecha futura → lo deja libre)."
)
h2("7.2 Importación masiva optimizada")
p(
    "La importación de estados (Enrolar Estado) está optimizada para minimizar viajes a "
    "la base de datos (importante porque PostgreSQL está en red). En vez de consultar y "
    "guardar fila por fila:"
)
nums([
    "Se parsea y deduplica todo el archivo en memoria (la última fila por MAC manda).",
    "Una sola consulta trae los televisores existentes (filter mac_address__in).",
    "Los televisores nuevos se crean con un bulk_create.",
    "Los Bloqueo nuevos (solo donde el estado cambia) se crean con bulk_create y los lock_status con bulk_update, todo dentro de una transacción.",
])
p(
    "Resultado: de ~5 consultas por fila a ~4 consultas en total. Tras importar, la "
    "aplicación detecta los televisores que cambiaron de estado y ofrece sincronizarlos."
)

# ----------------------------------------------------------------- 8
h1("8. Integración con el portal externo (Selenium)")
p(
    "Toda la automatización está en whaletv/portal_sync.py. El portal externo es una SPA "
    "(Vue + Element-UI), por lo que se controla como un usuario real; los selectores se "
    "basan en texto y atributos estables (no en los data-v-* dinámicos de Vue)."
)
h2("8.1 Construcción del navegador (_build_driver)")
bullets([
    "Chrome/Chromium con Selenium. En modo headless ('--headless=new').",
    "Flags pensados para contenedores y poca RAM: --no-sandbox, --disable-dev-shm-usage, --disable-gpu, --disable-extensions, imágenes desactivadas (imagesEnabled=false), etc.",
    "En el contenedor, el binario de Chromium y el driver se toman de las variables CHROME_BIN y CHROMEDRIVER. En local, Selenium Manager descarga el driver solo.",
])
h2("8.2 Flujo por televisor")
nums([
    "Login en el portal (_login): abre la URL de login, escribe correo/clave y espera la lista de dispositivos.",
    "Buscar el MAC (_abrir_detalle_por_mac): recorre las páginas de la lista de dispositivos hasta encontrar la MAC y abre su Detail.",
    "Leer el estado remoto (_leer_estado_remoto): interpreta el icono de candado (lock/unlock) del Detail.",
    "Comparar con el estado local (lock_status de la app).",
    "Dry-run (validación): no modifica nada; solo reporta y compara.",
    "Modo real (sincronización): Edit → fijar Lock/Unlock en el select de Element-UI → fijar la fecha en el date-picker → Save. Opcionalmente recarga para confirmar (configurable).",
])
p(
    "Funciones de entrada: sincronizar_televisor(televisor, dry_run) para una sola "
    "máquina (la usan 'Validar' y 'Sincronizar' del detalle), y ejecutar_job(...) para "
    "los trabajos masivos."
)

# ----------------------------------------------------------------- 9
h1("9. Procesos en segundo plano y paralelismo")
p(
    "Las operaciones masivas no bloquean la petición web: se ejecutan en un hilo aparte "
    "(threading.Thread daemon) que corre ejecutar_job(job_id, workers, dry_run)."
)
bullets([
    "Se crea un SyncJob y un SyncJobItem por televisor; los items pendientes se ponen en una cola compartida (queue.Queue).",
    "Cada worker es un hilo que abre UN navegador headless, hace UN login y va tomando televisores de la cola hasta vaciarla.",
    "Número de navegadores: 1 por televisor con un tope. Sincronización: tope configurable (WHALETV_PORTAL['MAX_WORKERS'], por defecto 6). Validación: tope fijo de 4 (1→1, 2→2, 3→3, 4 o más → 4).",
    "Solo puede haber un trabajo activo a la vez (comparten la sesión del portal). Hay auto-recuperación de trabajos 'colgados' (sin actividad por más de 10 minutos se cancelan).",
    "Cada worker cierra conexiones de BD al terminar (connections.close_all()) porque corre en su propio hilo.",
    "La página de progreso consulta un endpoint JSON (sync_progreso_api) por polling cada ~2 s y dibuja el avance; al terminar habilita la exportación.",
])
p(
    "Nota de despliegue: gunicorn corre con varios procesos (WEB_CONCURRENCY) y varios "
    "hilos por proceso (--threads 4, --timeout 120), lo que permite atender el polling "
    "mientras el hilo del trabajo sigue en segundo plano."
)

# ----------------------------------------------------------------- 10
h1("10. Importación y exportación de archivos")
bullets([
    "Lectura (importar): pandas lee .xlsx/.xls/.csv como texto; un mapeo de sinónimos tolera variaciones en los encabezados (p. ej. 'mac', 'mac address' → mac_address).",
    "Escritura (exportar y plantillas): openpyxl genera los .xlsx con encabezados con estilo y anchos de columna.",
    "Plantillas para importar: Enrolar Televisores (serial_number, mac_address, numero_credito) y Enrolar Estado (serial_number, mac_address, estado).",
    "Exportaciones: historial de sincronizaciones, Códigos Pin, resultado de un trabajo masivo y, en validación, solo los televisores cuyo estado NO coincide con el portal.",
])

# ----------------------------------------------------------------- 11
h1("11. Configuración y variables de entorno")
p(
    "La configuración vive en core/settings.py y se parametriza por variables de entorno "
    "(o un archivo .env en local)."
)
tabla(
    ["Variable", "Para qué"],
    [
        ["SECRET_KEY", "Clave secreta de Django (en Render se autogenera)."],
        ["DEBUG", "true/false. En producción debe ser false."],
        ["ALLOWED_HOSTS", "Hosts permitidos (Render añade RENDER_EXTERNAL_HOSTNAME automáticamente)."],
        ["CSRF_TRUSTED_ORIGINS", "Orígenes confiables para CSRF (se arman con los hosts https)."],
        ["DATABASE_URL", "Cadena de conexión a PostgreSQL (tiene prioridad; la parsea dj-database-url)."],
        ["DB_NAME/DB_USER/DB_PASSWORD/DB_HOST/DB_PORT", "Conexión a Postgres si no se usa DATABASE_URL."],
        ["TIME_ZONE", "Zona horaria (por defecto America/Bogota)."],
        ["WHALETV_PORTAL_LOGIN_URL / _DEVICE_LIST_URL", "URLs del portal externo."],
        ["WHALETV_PORTAL_EMAIL / _PASSWORD", "Credenciales del portal externo (no se versionan)."],
        ["WHALETV_PORTAL_HEADLESS", "true/false: navegador sin interfaz."],
        ["WHALETV_PORTAL_TIMEOUT", "Segundos de espera de Selenium (por defecto 30)."],
        ["WHALETV_PORTAL_MAX_WORKERS", "Tope de navegadores en paralelo (por defecto 6)."],
        ["WHALETV_PORTAL_CONFIRMAR_RECARGA", "Si recarga para confirmar el guardado (más lento)."],
        ["SECURE_SSL_REDIRECT / SECURE_HSTS_SECONDS", "Ajustes de seguridad en producción."],
    ],
)

# ----------------------------------------------------------------- 12
h1("12. Seguridad")
bullets([
    "Cuando DEBUG=false: SECURE_PROXY_SSL_HEADER (Render está detrás de un proxy HTTPS), SESSION_COOKIE_SECURE y CSRF_COOKIE_SECURE en True, redirección a HTTPS y HSTS opcional.",
    "Protección CSRF de Django en todos los formularios POST; CSRF_TRUSTED_ORIGINS configurado para el dominio público.",
    "Acceso restringido: todas las vistas internas requieren sesión iniciada (@login_required).",
    "Contraseñas validadas por los validadores estándar de Django y almacenadas con hash.",
    "Las credenciales del portal externo y la base de datos se entregan por variables de entorno (no quedan en el código).",
])

# ----------------------------------------------------------------- 13
h1("13. Base de datos y migraciones")
bullets([
    "Motor: PostgreSQL (django.db.backends.postgresql vía psycopg 3). Conexiones persistentes con CONN_MAX_AGE=60 y sslmode=require.",
    "En Render se enlaza la base y se inyecta DATABASE_URL (Internal Database URL), que tiene prioridad sobre las variables DB_*.",
    "Esquema gestionado con migraciones de Django (carpeta whaletv/migrations). Las migraciones se aplican en el despliegue (preDeployCommand / release) con 'python manage.py migrate'.",
    "Existe un db.sqlite3 en el repo que corresponde a pruebas locales antiguas; la configuración apunta a PostgreSQL.",
])

# ----------------------------------------------------------------- 14
h1("14. Despliegue en Render")
p(
    "La aplicación se despliega como servicio web en Render usando Docker, porque la "
    "automatización con Selenium necesita Chromium instalado (el runtime nativo de Python "
    "de Render no lo trae)."
)
h2("14.1 Dockerfile")
bullets([
    "Imagen base python:3.13-slim.",
    "Instala chromium, chromium-driver, libpq5 y fuentes; define CHROME_BIN=/usr/bin/chromium y CHROMEDRIVER=/usr/bin/chromedriver.",
    "Instala requirements, copia el código y ejecuta collectstatic.",
    "Arranca gunicorn: 'gunicorn core.wsgi:application --bind 0.0.0.0:$PORT --workers $WEB_CONCURRENCY --threads 4 --timeout 120'.",
])
h2("14.2 render.yaml (blueprint)")
bullets([
    "Servicio web tipo Docker, plan 'starter' (el free puede quedar corto con Chromium), región oregon.",
    "healthCheckPath: /login/.",
    "preDeployCommand: 'python manage.py migrate --noinput' (aplica migraciones antes de publicar).",
    "Variables: SECRET_KEY (autogenerada), DEBUG=false, WEB_CONCURRENCY=2, TIME_ZONE, y como secretos (sync:false) DATABASE_URL, WHALETV_PORTAL_EMAIL y WHALETV_PORTAL_PASSWORD.",
])
h2("14.3 Otros archivos de despliegue")
bullets([
    "Procfile: define 'release' (migrate) y 'web' (gunicorn) para el runtime nativo; útil como alternativa, pero Selenium NO funciona sin Chromium.",
    "build.sh: build para el runtime nativo (pip install, collectstatic, migrate). Incluye una advertencia de que Selenium requiere el despliegue con Docker.",
])

# ----------------------------------------------------------------- 15
h1("15. Rendimiento y límites a tener en cuenta")
bullets([
    "Cada navegador headless consume CPU y RAM; por eso el plan 'starter' o superior y el tope de navegadores en paralelo (MAX_WORKERS). En planes con poca RAM conviene bajarlo.",
    "El cuello de botella de la sincronización/validación es la navegación en el portal (buscar el MAC paginando). Mejora pendiente: usar el buscador del portal para no recorrer páginas.",
    "Se quitó la recarga de confirmación tras guardar (configurable) para acelerar; la confirmación se infiere de que reaparece el botón Edit.",
    "Las operaciones masivas corren en hilos; gunicorn usa --threads para no bloquear el polling de progreso.",
])

# ----------------------------------------------------------------- 16
h1("16. Ejecución en local (desarrollo)")
code(
    "python -m venv .venv && .venv\\Scripts\\activate   (Windows)\n"
    "pip install -r requirements.txt\n"
    "# Crear un archivo .env con SECRET_KEY, DEBUG=true, datos de la BD y\n"
    "# las variables WHALETV_PORTAL_* (ver core/settings.py).\n"
    "python manage.py migrate\n"
    "python manage.py createsuperuser\n"
    "python manage.py runserver   ->  http://127.0.0.1:8000/"
)
p(
    "En local, Selenium descarga el driver automáticamente y, si WHALETV_PORTAL_HEADLESS "
    "es false, se puede ver el navegador operando el portal."
)

doc.add_paragraph()
fin = doc.add_paragraph()
fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fin.add_run("— Fin de la documentación técnica —")
rf.font.color.rgb = GRIS

doc.save("Documentación técnica - Locking System.docx")
print("Documento técnico generado.")
