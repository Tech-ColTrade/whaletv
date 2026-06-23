# -*- coding: utf-8 -*-
"""Genera el manual de uso (Word) de la aplicación."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOMINIO = "https://whaletv.onrender.com"
WHALE = RGBColor(0xF6, 0x18, 0x6A)
GRIS = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Estilo base
base = doc.styles['Normal']
base.font.name = 'Calibri'
base.font.size = Pt(11)


def h1(texto):
    p = doc.add_heading(texto, level=1)
    for r in p.runs:
        r.font.color.rgb = WHALE
    return p


def h2(texto):
    return doc.add_heading(texto, level=2)


def parrafo(texto, negrita=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = negrita
    return p


def url(texto, ruta):
    p = doc.add_paragraph()
    r = p.add_run(texto + "  ")
    r.bold = True
    r2 = p.add_run(DOMINIO + ruta)
    r2.font.color.rgb = WHALE
    return p


def vinneta(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def numerada(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def tabla(encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, e in enumerate(encabezados):
        hdr[i].text = e
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = str(val)
    doc.add_paragraph()
    return t


# ---------------------------------------------------------------------------
# PORTADA
# ---------------------------------------------------------------------------
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = titulo.add_run("Locking System")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = WHALE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Manual de uso")
rs.font.size = Pt(16)
rs.font.color.rgb = GRIS

dom = doc.add_paragraph()
dom.alignment = WD_ALIGN_PARAGRAPH.CENTER
dom.add_run(DOMINIO)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. QUÉ ES
# ---------------------------------------------------------------------------
h1("1. ¿Qué es esta aplicación y qué es el portal de bloqueo?")

parrafo(
    "Esta aplicación es tu panel de control para administrar los televisores que se "
    "venden a cuotas. Aquí registras cada televisor, llevas su número de crédito y "
    "decides si debe quedar BLOQUEADO o DESBLOQUEADO. Es donde tú trabajas el día a día."
)
parrafo(
    "El portal de bloqueo es un sistema externo (aparte de esta aplicación) que es el "
    "que realmente bloquea o desbloquea los televisores. Es el sistema oficial donde "
    "queda registrado el estado verdadero de cada equipo."
)

h2("¿Cómo trabajan juntas?")
parrafo(
    "Tú defines el estado de cada televisor en esta aplicación (por ejemplo, marcar una "
    "MAC como Bloqueada). Ese estado todavía no llega solo al portal: la aplicación debe "
    "“sincronizar”, es decir, entrar al portal de bloqueo de forma automática y aplicar "
    "allí el mismo estado que tú pusiste. A partir de ese momento el televisor queda "
    "bloqueado o desbloqueado de verdad."
)
parrafo("En resumen:")
vinneta([
    "Aplicación = donde tú decides y organizas el estado de los televisores.",
    "Portal de bloqueo = donde el estado se aplica de verdad sobre el equipo.",
    "Sincronizar = pasar el estado de la aplicación al portal de bloqueo.",
    "Validar = consultar el estado real en el portal y compararlo con lo que dice la aplicación.",
])

# ---------------------------------------------------------------------------
# 2. CONCEPTOS
# ---------------------------------------------------------------------------
h1("2. Palabras que vas a ver en toda la aplicación")
tabla(
    ["Palabra", "Qué significa"],
    [
        ["Dirección MAC", "El código que identifica a cada televisor (algo como B4:04:29:7E:3A:ED)."],
        ["Número de serie", "El serial del televisor."],
        ["N° Crédito", "El número del crédito asociado al televisor (solo números, hasta 60 dígitos)."],
        ["Estado", "Si el televisor está Bloqueado o Desbloqueado."],
        ["Bloqueado", "El televisor no se puede usar (está restringido)."],
        ["Desbloqueado", "El televisor funciona normalmente."],
        ["Sincronizar", "Aplicar en el portal de bloqueo el estado que tiene la aplicación."],
        ["Validar", "Consultar el estado real en el portal y compararlo con la aplicación."],
        ["Código de Acceso", "Una clave que te da el sistema para generar un Código Pin de desbloqueo."],
        ["Código Pin", "El número que desbloquea manualmente un televisor."],
        ["Fecha portal", "La fecha que la aplicación pone en el portal para mantener el estado (ver más abajo)."],
    ],
)
parrafo("Sobre la “Fecha portal”:", negrita=True)
parrafo(
    "Cuando un televisor queda Bloqueado, la aplicación coloca en el portal una fecha de "
    "30 días hacia atrás (una fecha ya vencida), lo que hace que el portal lo mantenga "
    "bloqueado. Cuando queda Desbloqueado, coloca una fecha de 30 días hacia adelante "
    "(una fecha futura), para que el portal lo deje libre. Esto se calcula solo; tú no "
    "tienes que escribir ninguna fecha."
)

# ---------------------------------------------------------------------------
# 3. INICIAR SESIÓN
# ---------------------------------------------------------------------------
h1("3. Iniciar sesión")
url("Página:", "/login/")
parrafo(
    "Es la primera pantalla. Escribe tu Correo y tu Contraseña y presiona el botón "
    "“Iniciar sesión”. Si los datos son correctos entrarás a la aplicación; si no, "
    "aparecerá un aviso de “Correo o contraseña incorrectos”."
)
parrafo(
    "Mientras no hayas iniciado sesión, la aplicación no te dejará entrar a ninguna "
    "otra pantalla."
)

# ---------------------------------------------------------------------------
# 4. INICIO
# ---------------------------------------------------------------------------
h1("4. Página de inicio")
url("Página:", "/")
parrafo(
    "Después de iniciar sesión ves un saludo de bienvenida y tu correo. Tiene un botón "
    "“Gestionar televisores →” que te lleva al listado de televisores, que es el corazón "
    "de la aplicación."
)

# ---------------------------------------------------------------------------
# 5. LISTADO DE TELEVISORES
# ---------------------------------------------------------------------------
h1("5. Listado de televisores (pantalla principal)")
url("Página:", "/televisores/")
parrafo(
    "Aquí ves todos los televisores registrados en una tabla. Arriba hay un buscador y "
    "una fila de botones; abajo, la tabla."
)

h2("El buscador")
parrafo(
    "El cuadro “Buscar MAC, serial, crédito, nombre...” te permite filtrar la tabla. "
    "Escribe lo que quieras encontrar (una MAC, un serial o un número de crédito) y "
    "presiona “Buscar”."
)

h2("Las columnas de la tabla")
tabla(
    ["Columna", "Qué muestra"],
    [
        ["Dirección MAC", "El identificador del televisor. Es un enlace: si haces clic, abres el detalle de ese televisor."],
        ["Número de serie", "El serial del televisor."],
        ["N° Crédito", "El número de crédito."],
        ["Estado", "Una etiqueta: Bloqueado (rojo) o Desbloqueado (verde)."],
        ["Acciones", "Los botones para ese televisor (ver abajo)."],
    ],
)

h2("Botones de cada fila (columna Acciones)")
vinneta([
    "Estado: te lleva a la pantalla para cambiar el estado de ese televisor (Bloqueado/Desbloqueado), sin tener que entrar al detalle.",
    "Editar: abre el formulario para corregir la MAC, el serial o el número de crédito.",
    "Eliminar: borra el televisor (pide confirmación).",
])

h2("Botones de arriba (acciones generales)")
parrafo("De izquierda a derecha:")
vinneta([
    "Validación masiva: revisa en el portal de bloqueo el estado real de TODOS los televisores y te dice cuáles NO coinciden con la aplicación. NO cambia nada. (Se explica en el punto 11.)",
    "⟳ Sincronizar todo: aplica en el portal de bloqueo el estado de TODOS los televisores. SÍ cambia el portal. Pide confirmación. (Punto 10.)",
    "Enrolar Televisores: cargar muchos televisores de una vez con un archivo de Excel. (Punto 8.)",
    "Enrolar Estado: cargar el estado (Bloqueado/Desbloqueado) de muchos televisores de una vez con un archivo de Excel. (Punto 9.)",
    "⬇ Exportar sincronizaciones: descarga un Excel con todo el historial de sincronizaciones. (Punto 12.)",
    "⬇ Exportar Códigos Pin: descarga un Excel con todos los Códigos Pin generados. (Punto 13.)",
    "+ Nuevo: crea un televisor a mano, uno por uno. (Punto 6.)",
])

# ---------------------------------------------------------------------------
# 6. CREAR / EDITAR / ELIMINAR
# ---------------------------------------------------------------------------
h1("6. Crear, editar o eliminar un televisor (uno por uno)")
url("Crear:", "/televisores/nuevo/")
parrafo(
    "Con el botón “+ Nuevo” llegas a un formulario con tres campos: Dirección MAC, "
    "Número de serie y N° Crédito. El número de crédito solo admite dígitos. Presiona "
    "“Guardar” y el televisor queda registrado (inicia como Desbloqueado)."
)
url("Editar:", "/televisores/<número>/editar/")
parrafo(
    "Desde el botón “Editar” de la fila. Cambias los datos y guardas. (El estado no se "
    "cambia aquí, sino en la pantalla de Estado.)"
)
url("Eliminar:", "/televisores/<número>/eliminar/")
parrafo(
    "Desde el botón “Eliminar”. Muestra los datos del televisor y pide confirmación "
    "antes de borrarlo. Esta acción no se puede deshacer."
)

# ---------------------------------------------------------------------------
# 7. CAMBIAR ESTADO
# ---------------------------------------------------------------------------
h1("7. Cambiar el estado de un televisor")
url("Página:", "/televisores/<número>/historico/")
parrafo(
    "Llegas con el botón “Estado” de la tabla. Esta pantalla muestra el estado actual "
    "del televisor y el historial de cambios."
)
vinneta([
    "Para cambiar el estado: en “Registrar estado” eliges Bloqueado o Desbloqueado y presionas “Registrar”. El televisor queda con ese estado en la aplicación.",
    "Abajo ves el historial: la fecha de cada cambio, el estado y el serial. Cada registro se puede eliminar.",
])
parrafo(
    "Importante: cambiar el estado aquí lo deja listo en la aplicación. Para que ese "
    "estado llegue al portal de bloqueo todavía hay que Sincronizar (botón “Sincronizar” "
    "del detalle, o “Sincronizar todo” del listado)."
)

# ---------------------------------------------------------------------------
# 8. ENROLAR TELEVISORES
# ---------------------------------------------------------------------------
h1("8. Enrolar Televisores (cargar muchos con Excel)")
url("Página:", "/televisores/importar/")
parrafo(
    "Sirve para registrar muchos televisores de una sola vez subiendo un archivo de "
    "Excel. Si un televisor ya existe (misma MAC) se actualiza; si no existe, se crea."
)
parrafo("Pasos:")
numerada([
    "Presiona “Descargar plantilla Excel” para obtener el archivo de ejemplo ya con las columnas correctas.",
    "Llena la plantilla con tus televisores.",
    "Elige el archivo y presiona “Importar”.",
])
parrafo("Columnas de la plantilla de Enrolar Televisores:", negrita=True)
tabla(
    ["Columna", "Ejemplo", "¿Obligatoria?"],
    [
        ["serial_number", "B4:04:29:7E:3A:AA", "No"],
        ["mac_address", "B4:04:29:7E:3A:AA", "Sí (identifica el televisor)"],
        ["numero_credito", "1234567890", "No (solo números, hasta 60 dígitos)"],
    ],
)
parrafo(
    "Al terminar, la aplicación te muestra cuántos televisores se crearon y cuántos se "
    "actualizaron."
)

# ---------------------------------------------------------------------------
# 9. ENROLAR ESTADO
# ---------------------------------------------------------------------------
h1("9. Enrolar Estado (cargar el estado de muchos con Excel)")
url("Página:", "/bloqueos/importar/")
parrafo(
    "Sirve para fijar de una sola vez el estado (Bloqueado/Desbloqueado) de muchos "
    "televisores subiendo un Excel. Si la MAC no existe, se crea el televisor con ese "
    "estado."
)
parrafo("Pasos:")
numerada([
    "Presiona “Descargar plantilla Excel”.",
    "Llena la columna “estado” con la palabra bloqueado o desbloqueado en cada televisor.",
    "Elige el archivo y presiona “Importar”.",
])
parrafo("Columnas de la plantilla de Enrolar Estado:", negrita=True)
tabla(
    ["Columna", "Ejemplo", "¿Obligatoria?"],
    [
        ["serial_number", "B4:04:29:7E:3A:EE", "No"],
        ["mac_address", "B4:04:29:7E:3A:EE", "Sí (identifica el televisor)"],
        ["estado", "bloqueado / desbloqueado", "Sí"],
    ],
)
h2("Qué pasa después de importar el estado")
parrafo(
    "La aplicación detecta cuáles televisores CAMBIARON de estado y te lo muestra en un "
    "aviso: “Se detectaron cambios de estado”. Ahí puedes:"
)
vinneta([
    "Previsualizar cambios: ver una tabla con cada televisor que cambió, mostrando el cambio (por ejemplo, Desbloqueado → Bloqueado) y la fecha que se pondrá en el portal.",
    "Aceptar y sincronizar: aplicar esos cambios en el portal de bloqueo de una vez.",
    "Salir sin sincronizar: dejar los cambios guardados en la aplicación y sincronizar después.",
])

# ---------------------------------------------------------------------------
# 10. SINCRONIZAR
# ---------------------------------------------------------------------------
h1("10. Sincronizar (aplicar el estado en el portal de bloqueo)")
parrafo(
    "Sincronizar es el momento en que la aplicación entra al portal de bloqueo y aplica "
    "allí el estado que tú definiste. Hay dos formas:"
)
h2("Sincronizar un solo televisor")
url("Desde el detalle:", "/detailtv/<MAC>/")
parrafo(
    "En el detalle del televisor, el botón “Sincronizar” aplica en el portal el estado "
    "(Bloqueado/Desbloqueado) y la fecha de ese televisor. Pide confirmación porque sí "
    "cambia el portal."
)
h2("Sincronizar todos (Sincronización masiva)")
url("Botón en:", "/televisores/")
parrafo(
    "El botón “⟳ Sincronizar todo” aplica el estado de TODOS los televisores en el "
    "portal. La aplicación abre varios navegadores automáticos a la vez para ir más "
    "rápido (un navegador por televisor, hasta un máximo de varios en paralelo)."
)
parrafo("Mientras corre, ves una pantalla de progreso con:")
vinneta([
    "Un círculo con el porcentaje y cuántos televisores van procesados.",
    "Aplicados: a cuántos se les cambió el estado en el portal.",
    "Sin cambios: cuántos ya estaban igual en el portal.",
    "Con error: cuántos fallaron.",
    "Al terminar aparece un botón “⬇ Exportar a Excel” con el detalle de esa sincronización.",
])
parrafo(
    "Puedes dejar esa pantalla abierta hasta que termine, o cancelar el proceso con el "
    "botón “Cancelar”."
)

# ---------------------------------------------------------------------------
# 11. VALIDAR
# ---------------------------------------------------------------------------
h1("11. Validar (consultar el estado real del portal)")
parrafo(
    "Validar NO cambia nada. Solo entra al portal de bloqueo, lee el estado real de los "
    "televisores y lo compara con lo que dice la aplicación. Sirve para verificar que "
    "todo esté igual."
)
h2("Validar un solo televisor")
parrafo(
    "En el detalle del televisor, el botón “Validar” consulta el portal y te dice, por "
    "ejemplo: “El televisor … está Bloqueado en el portal, igual que en la app” o "
    "“… está Desbloqueado en el portal, pero Bloqueado en la app”."
)
h2("Validación masiva")
url("Botón en:", "/televisores/")
parrafo(
    "El botón “Validación masiva” hace lo mismo pero con todos los televisores a la vez "
    "(un navegador por televisor, hasta 4 a la vez). Es como el “Validar” individual pero "
    "para todos."
)
parrafo("En la pantalla de progreso verás:")
vinneta([
    "Coinciden: cuántos televisores están igual en la aplicación y en el portal.",
    "No coinciden: cuántos están diferentes.",
    "Con error: cuántos no se pudieron revisar.",
    "Al terminar aparece “⬇ Exportar los que no coinciden”, que descarga un Excel solo con los televisores diferentes, mostrando el Estado en el portal y el Estado en la app.",
])
parrafo(
    "Nota: solo se puede correr un proceso a la vez (una sincronización o una "
    "validación). Si ya hay uno en curso, la aplicación te lleva a su pantalla de progreso."
)

# ---------------------------------------------------------------------------
# 12. SINCRONIZACIONES (bitácora)
# ---------------------------------------------------------------------------
h1("12. Historial de sincronizaciones")
url("Página:", "/sincronizaciones/")
parrafo(
    "Es la bitácora: muestra cada vez que se sincronizó un televisor, quién lo hizo y "
    "cuándo. Tiene buscador y un botón “⬇ Exportar todo a Excel”."
)
parrafo("La tabla y el Excel incluyen:")
vinneta([
    "Fecha y quién sincronizó.",
    "Dirección MAC del televisor.",
    "Estado anterior y Estado final (de qué estado venía y a cuál pasó).",
    "Resultado: Aplicado (se cambió) o Sin cambios (ya estaba igual).",
    "Tipo: individual o masivo.",
])

# ---------------------------------------------------------------------------
# 13. CÓDIGOS PIN
# ---------------------------------------------------------------------------
h1("13. Códigos Pin y Desbloqueo manual")
url("Página:", "/pincodes/")
parrafo(
    "El Desbloqueo manual sirve para generar un Código Pin que desbloquea un televisor "
    "directamente. Se hace desde el detalle del televisor."
)
parrafo("Pasos para generar un Código Pin:")
numerada([
    "Entra al detalle del televisor (clic en la MAC) y presiona “Generar Código Pin”.",
    "Escribe el Código de Acceso y presiona “Continuar”.",
    "La aplicación se conecta al portal y te muestra el Código Pin. Lo puedes copiar.",
    "Con “Listo y sincronizar” además aplica el estado y la fecha en el portal.",
])
parrafo(
    "Todos los Códigos Pin generados quedan guardados. En la página /pincodes/ los ves "
    "todos, con buscador y botón “⬇ Exportar todo a Excel”. El Excel incluye Fecha, "
    "Dirección MAC, Código de Acceso, Código Pin y quién lo generó."
)

# ---------------------------------------------------------------------------
# 14. DETALLE
# ---------------------------------------------------------------------------
h1("14. Detalle de un televisor")
url("Página:", "/detailtv/<MAC>/")
parrafo(
    "Se abre al hacer clic en la MAC dentro del listado. Reúne todo lo de un televisor:"
)
vinneta([
    "Información del dispositivo: Dirección MAC, Número de serie, N° Crédito, Estado y Fecha portal, con los botones “Validar” y “Sincronizar”.",
    "Desbloqueo manual: el botón “Generar Código Pin”.",
    "Registros: accesos rápidos a las Sincronizaciones, los Bloqueos (historial de estado) y los Códigos Pin de ese televisor, cada uno con su cantidad.",
])

# ---------------------------------------------------------------------------
# 15. RESUMEN DE ARCHIVOS
# ---------------------------------------------------------------------------
h1("15. Resumen de las plantillas y los archivos que se descargan")
h2("Plantillas para subir (importar)")
tabla(
    ["Para qué", "Columnas del Excel"],
    [
        ["Enrolar Televisores", "serial_number, mac_address, numero_credito"],
        ["Enrolar Estado", "serial_number, mac_address, estado (bloqueado/desbloqueado)"],
    ],
)
h2("Archivos que se descargan (exportar)")
tabla(
    ["Botón / pantalla", "Qué trae el Excel"],
    [
        ["Exportar sincronizaciones", "Fecha, quién, MAC, Estado anterior, Estado final, Resultado, Tipo"],
        ["Exportar Códigos Pin", "Fecha, MAC, Código de Acceso, Código Pin, quién lo generó"],
        ["Exportar (al terminar una sincronización)", "MAC, serial, N° Crédito, Estado anterior, Estado final, resultado, mensaje"],
        ["Exportar los que no coinciden (validación)", "MAC, serial, N° Crédito, Estado en el portal, Estado en la app"],
    ],
)

# ---------------------------------------------------------------------------
# 16. FLUJO RECOMENDADO
# ---------------------------------------------------------------------------
h1("16. Un flujo de trabajo recomendado")
numerada([
    "Registra los televisores: a mano con “+ Nuevo”, o muchos a la vez con “Enrolar Televisores”.",
    "Define el estado: a mano con el botón “Estado” de cada fila, o muchos a la vez con “Enrolar Estado”.",
    "Sincroniza: usa “⟳ Sincronizar todo” (o “Sincronizar” en un televisor) para aplicar los estados en el portal de bloqueo.",
    "Verifica: usa “Validación masiva” para confirmar que el portal quedó igual que la aplicación, y exporta los que no coinciden si hace falta.",
    "Consulta el historial: en “Historial de sincronizaciones” y en “Códigos Pin” cuando lo necesites.",
])

doc.add_paragraph()
fin = doc.add_paragraph()
fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fin.add_run("— Fin del manual —")
rf.font.color.rgb = GRIS

doc.save("Manual de uso - Locking System.docx")
print("Documento generado.")
